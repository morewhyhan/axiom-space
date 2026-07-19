#!/usr/bin/env python3
"""Build the AXIOM Space 2026 Software Cup A3 submission DOCX set.

The source of truth remains docs/04-提交文档/*.md.  This builder creates an
external-review Word edition without rewriting or shortening the Markdown.

Resolved design system
----------------------
Base preset: ``compact_reference_guide``.
First-page pattern: ``editorial_cover``.

Named suite-wide overrides (applied consistently to every document):

* ``A4_CN_SUBMISSION``: A4 portrait, 22 mm side margins, 20 mm top/bottom,
  10 mm header/footer distance, 166 mm usable width.
* ``CN_TECH_TYPE``: SimSun 10.5 pt body; Microsoft YaHei headings; Consolas
  code; 1.25 body line spacing.
* ``AXIOM_BRAND``: navy #10243E, teal #14A89A, harvest gold #D6A64B,
  ink #172235, blue-gray #E9EFF4, light gray #F5F7F9.
* ``DENSE_EVIDENCE_TABLE``: fixed-DXA full-width tables, explicit grid and
  cell widths, compact 7.0-8.5 pt table text, repeating header rows.

The output is intentionally a set of independent DOCX files.  Cross-document
Markdown links are rewritten to the sibling DOCX files; local figures and all
Mermaid diagrams are embedded so the Word package is self-contained.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import math
import os
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from collections import defaultdict, deque
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = REPO_ROOT / "docs" / "04-提交文档"
OUTPUT_DIR = SOURCE_DIR / "DOCX提交版"
WORK_DIR = REPO_ROOT / "tmp" / "submission-docx"
DIAGRAM_DIR = WORK_DIR / "diagrams"
HERO_DIR = WORK_DIR / "covers"

TEAM_NAME = "给春稗以秋实"
COMPETITION = "2026 年软件杯 · A3 组"
PRODUCT = "AXIOM Space"
VERSION = "提交版 · 2026.07"

# A4_CN_SUBMISSION token map.
PAGE_W_MM = 210
PAGE_H_MM = 297
MARGIN_SIDE_MM = 22
MARGIN_TOP_MM = 20
MARGIN_BOTTOM_MM = 20
HEADER_MM = 10
FOOTER_MM = 10
CONTENT_MM = PAGE_W_MM - MARGIN_SIDE_MM * 2
CONTENT_DXA = round(CONTENT_MM / 25.4 * 1440)
TABLE_INDENT_DXA = 120
TABLE_WIDTH_DXA = CONTENT_DXA - TABLE_INDENT_DXA
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_SIDE = 120

NAVY = "10243E"
NAVY_2 = "183B5B"
TEAL = "14A89A"
TEAL_DARK = "0F756D"
GOLD = "D6A64B"
INK = "172235"
MUTED = "657487"
BLUE_GRAY = "E9EFF4"
LIGHT_GRAY = "F5F7F9"
WHITE = "FFFFFF"
LINE = "D6DEE6"
CODE_BG = "F3F6F8"
CALLOUT_BG = "ECF8F6"
RISK_BG = "FFF5E5"

BODY_FONT_CN = "宋体"
BODY_FONT_LATIN = "Arial"
HEADING_FONT_CN = "微软雅黑"
HEADING_FONT_LATIN = "Microsoft YaHei"
CODE_FONT = "Consolas"

FONT_REGULAR = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")

DOC_INFO = {
    "00": {
        "title": "参赛提交文档索引",
        "english": "Submission Document Index",
        "purpose": "定义整套参赛材料的职责边界、证据关系与推荐阅读路径。",
        "toc_depth": 2,
    },
    "01": {
        "title": "用户研究报告",
        "english": "User Research Report",
        "purpose": "以公开研究、行业数据与用户反馈还原新时代大学生的真实学习需求。",
        "toc_depth": 1,
    },
    "02": {
        "title": "软件需求规格说明书",
        "english": "Software Requirements Specification",
        "purpose": "将用户问题与赛题要求转化为可追溯、可验证、可验收的软件需求。",
        "toc_depth": 1,
    },
    "03": {
        "title": "系统设计与开发说明书",
        "english": "System Design & Development Specification",
        "purpose": "完整说明系统架构、智能体设计、核心实现、系统集成、优化与工程边界。",
        "toc_depth": 1,
    },
    "04": {
        "title": "测试说明书与测试报告",
        "english": "Test Specification & Test Report",
        "purpose": "统一呈现测试策略、用例体系、执行证据、异常定位与验收结论。",
        "toc_depth": 1,
    },
    "05": {
        "title": "部署说明书",
        "english": "Deployment Guide",
        "purpose": "说明安装、配置、启动、健康检查、故障恢复与演示环境复现方法。",
        "toc_depth": 1,
    },
    "06": {
        "title": "用户使用说明书",
        "english": "User Manual",
        "purpose": "指导评委与学习者完成 AXIOM Space 的核心学习闭环与关键操作。",
        "toc_depth": 1,
    },
    "07": {
        "title": "《软件设计模式》系统化学习资料",
        "english": "Systematic Learning Guide to Software Design Patterns",
        "purpose": "提供可直接学习、授课、导入系统并开展掌握评估的完整课程资料。",
        "toc_depth": 2,
    },
    "08": {
        "title": "开源组件使用与合规说明",
        "english": "Open-source Components & Compliance",
        "purpose": "披露第三方软件、容器、模型、用途、许可证义务与团队自研边界。",
        "toc_depth": 1,
    },
    "09": {
        "title": "团队级 AI Coding 方法论\n与 AI 辅助开发流程说明",
        "english": "Team-level AI Coding Methodology & Development Workflow",
        "purpose": "完整呈现团队研究并实践的 Rules、Specification、Skills、Loop 与 Harness 方法论。",
        "toc_depth": 2,
    },
    "10": {
        "title": "产品创新与设计决策全记录",
        "english": "Product Innovation & Design Decision Record",
        "purpose": "完整保留从问题识别、方案比较、约束推导到产品与技术决策落地的证据链。",
        "toc_depth": 1,
    },
}


@dataclass
class HeadingInfo:
    source_level: int
    word_level: int
    text: str
    bookmark: str
    source_slug: str


@dataclass
class BuildStats:
    code: str
    source_lines: int = 0
    source_chars: int = 0
    paragraphs: int = 0
    headings: int = 0
    tables: int = 0
    table_rows: int = 0
    figures: int = 0
    mermaid: int = 0
    code_blocks: int = 0
    hyperlinks: int = 0
    missing_assets: int = 0
    output_bytes: int = 0


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    latin: str = BODY_FONT_LATIN,
    east_asia: str = BODY_FONT_CN,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(
    cell,
    *,
    top: int = CELL_MARGIN_TOP_BOTTOM,
    start: int = CELL_MARGIN_SIDE,
    bottom: int = CELL_MARGIN_TOP_BOTTOM,
    end: int = CELL_MARGIN_SIDE,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = LINE, size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if not widths_dxa:
        return
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.width = Inches(width / 1440)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_table_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, *, side: str, color: str, size: int = 16, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def configure_section(section, *, cover: bool = False) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    if cover:
        section.top_margin = Mm(16)
        section.bottom_margin = Mm(16)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
    else:
        section.top_margin = Mm(MARGIN_TOP_MM)
        section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
        section.left_margin = Mm(MARGIN_SIDE_MM)
        section.right_margin = Mm(MARGIN_SIDE_MM)
    section.header_distance = Mm(HEADER_MM)
    section.footer_distance = Mm(FOOTER_MM)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT_LATIN
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_CN)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        1: (17.0, NAVY, 18, 9, False),
        2: (13.5, TEAL_DARK, 13, 6, False),
        3: (11.5, NAVY_2, 10, 5, False),
        4: (10.5, TEAL_DARK, 8, 4, False),
        5: (10.0, MUTED, 7, 3, False),
    }
    for level, (size, color, before, after, page_break) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = HEADING_FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), HEADING_FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), HEADING_FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT_CN)
        sf = style.paragraph_format
        sf.space_before = Pt(before)
        sf.space_after = Pt(after)
        sf.line_spacing = 1.15
        sf.keep_with_next = True
        sf.keep_together = True
        sf.page_break_before = page_break

    if "AXIOM Body" not in [s.name for s in styles]:
        body = styles.add_style("AXIOM Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        body = styles["AXIOM Body"]
    body.base_style = normal
    body.paragraph_format.first_line_indent = Pt(21)

    if "AXIOM Quote" not in [s.name for s in styles]:
        quote = styles.add_style("AXIOM Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["AXIOM Quote"]
    quote.base_style = normal
    quote.paragraph_format.left_indent = Mm(5)
    quote.paragraph_format.right_indent = Mm(3)
    quote.paragraph_format.space_before = Pt(5)
    quote.paragraph_format.space_after = Pt(7)
    quote.paragraph_format.line_spacing = 1.2

    if "AXIOM Code" not in [s.name for s in styles]:
        code = styles.add_style("AXIOM Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["AXIOM Code"]
    code.font.name = CODE_FONT
    code.font.size = Pt(8.0)
    code.font.color.rgb = rgb(INK)
    code._element.rPr.rFonts.set(qn("w:ascii"), CODE_FONT)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), CODE_FONT)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT_CN)
    code.paragraph_format.left_indent = Mm(4)
    code.paragraph_format.right_indent = Mm(2)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = HEADING_FONT_LATIN
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = rgb(MUTED)
    caption._element.rPr.rFonts.set(qn("w:ascii"), HEADING_FONT_LATIN)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), HEADING_FONT_LATIN)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT_CN)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), BODY_FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), BODY_FONT_CN)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.5, color=MUTED)


def set_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_header_footer(section, code: str, title: str) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    table = header.add_table(rows=1, cols=2, width=Mm(CONTENT_MM))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [round(CONTENT_DXA * 0.56), CONTENT_DXA - round(CONTENT_DXA * 0.56)]
    set_table_geometry(table, widths, indent_dxa=0)
    set_table_no_borders(table)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(f"{PRODUCT.upper()}  |  {COMPETITION}")
    set_run_font(lr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.0, color=MUTED, bold=True)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(f"[{code}] {title.replace(chr(10), ' ')}")
    set_run_font(rr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.0, color=MUTED)

    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_after = Pt(0)
    ft = footer.add_table(rows=1, cols=3, width=Mm(CONTENT_MM))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    ft.autofit = False
    fwidths = [round(CONTENT_DXA * 0.42), round(CONTENT_DXA * 0.32), CONTENT_DXA - round(CONTENT_DXA * 0.74)]
    set_table_geometry(ft, fwidths, indent_dxa=0)
    set_table_no_borders(ft)
    for cell in ft.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
    p0 = ft.rows[0].cells[0].paragraphs[0]
    r0 = p0.add_run(TEAM_NAME)
    set_run_font(r0, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.0, color=MUTED)
    p1 = ft.rows[0].cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"DOC {code}")
    set_run_font(r1, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.0, color=TEAL_DARK, bold=True)
    p2 = ft.rows[0].cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("PAGE ")
    set_run_font(r2, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.0, color=MUTED)
    add_page_field(p2)


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(str(path), size)
    except Exception:
        return ImageFont.load_default()


def make_cover_hero(code: str) -> Path:
    HERO_DIR.mkdir(parents=True, exist_ok=True)
    out = HERO_DIR / f"cover-{code}.png"
    if out.exists():
        return out

    width, height = 2200, 820
    base = Image.new("RGB", (width, height), f"#{NAVY}")
    px = base.load()
    c1 = tuple(int(NAVY[i : i + 2], 16) for i in (0, 2, 4))
    c2 = tuple(int(NAVY_2[i : i + 2], 16) for i in (0, 2, 4))
    for y in range(height):
        fy = y / max(1, height - 1)
        for x in range(width):
            fx = x / max(1, width - 1)
            mix = min(1.0, 0.22 + 0.62 * fx + 0.16 * fy)
            px[x, y] = tuple(round(c1[k] * (1 - mix) + c2[k] * mix) for k in range(3))

    glow = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    gd = ImageDraw.Draw(glow)
    gd.ellipse((1400, -440, 2480, 640), fill=(20, 168, 154, 72))
    glow = glow.filter(ImageFilter.GaussianBlur(92))
    base = Image.alpha_composite(base.convert("RGBA"), glow)
    draw = ImageDraw.Draw(base)

    grid_color = (255, 255, 255, 18)
    for x in range(0, width, 110):
        draw.line((x, 0, x, height), fill=grid_color, width=1)
    for y in range(0, height, 110):
        draw.line((0, y, width, y), fill=grid_color, width=1)

    draw.arc((1280, -430, 2500, 790), 135, 318, fill=(20, 168, 154, 205), width=8)
    draw.arc((1460, -260, 2310, 590), 120, 328, fill=(214, 166, 75, 210), width=5)
    draw.arc((1540, -100, 2150, 510), 102, 335, fill=(255, 255, 255, 92), width=3)
    draw.ellipse((1840, 138, 1882, 180), fill=(214, 166, 75, 255))
    draw.ellipse((2020, 370, 2046, 396), fill=(20, 168, 154, 255))

    font_brand = load_font(FONT_BOLD, 116)
    font_sub = load_font(FONT_REGULAR, 34)
    font_doc = load_font(FONT_BOLD, 68)
    font_num = load_font(FONT_BOLD, 210)
    draw.text((150, 176), "AXIOM", font=font_brand, fill=(255, 255, 255, 255))
    draw.text((156, 298), "SPACE", font=font_brand, fill=(255, 255, 255, 255))
    draw.rounded_rectangle((154, 455, 870, 518), radius=22, fill=(20, 168, 154, 230))
    draw.text((182, 468), "INTELLIGENT LEARNING AGENT SYSTEM", font=font_sub, fill=(255, 255, 255, 255))

    draw.text((1638, 500), f"DOC {code}", font=font_doc, fill=(255, 255, 255, 210))
    draw.text((1630, 225), code, font=font_num, fill=(255, 255, 255, 245))
    draw.rectangle((145, 650, 680, 666), fill=(214, 166, 75, 255))
    draw.rectangle((700, 650, 910, 666), fill=(20, 168, 154, 255))

    base.convert("RGB").save(out, quality=96)
    return out


def add_cover(doc: Document, code: str, info: dict[str, object]) -> None:
    section = doc.sections[0]
    configure_section(section, cover=True)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run()
    run.add_picture(str(make_cover_hero(code)), width=Mm(170))
    p.paragraph_format.keep_with_next = True

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(10)
    kr = kicker.add_run(COMPETITION)
    set_run_font(kr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=11.5, color=GOLD, bold=True)
    kr.font.all_caps = True

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(8)
    tr = title_p.add_run(str(info["title"]))
    set_run_font(tr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=27.5, color=NAVY, bold=True)

    en = doc.add_paragraph()
    en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en.paragraph_format.space_after = Pt(18)
    er = en.add_run(str(info["english"]))
    set_run_font(er, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=11.0, color=MUTED, italic=True)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Mm(12)
    summary.paragraph_format.right_indent = Mm(12)
    summary.paragraph_format.space_after = Pt(20)
    sr = summary.add_run(str(info["purpose"]))
    set_run_font(sr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=10.5, color=INK)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(14)
    set_paragraph_border(line, side="bottom", color=GOLD, size=18, space=4)

    meta = [
        f"项目：{PRODUCT}",
        f"参赛队伍：{TEAM_NAME}",
        f"文档编号：{code}",
        f"版本：{VERSION}",
    ]
    for item in meta:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_after = Pt(3)
        mr = mp.add_run(item)
        set_run_font(mr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=9.8, color=MUTED, bold=item.startswith("参赛队伍"))


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, *, color: str = TEAL_DARK, bold: bool = False) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), HEADING_FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), HEADING_FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), HEADING_FONT_CN)
    rpr.append(rfonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    if bold:
        rpr.append(OxmlElement("w:b"))
    no_u = OxmlElement("w:u")
    no_u.set(qn("w:val"), "none")
    rpr.append(no_u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_hyperlink(paragraph, text: str, target: str, *, bold: bool = False, italic: bool = False) -> None:
    rid = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), HEADING_FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), HEADING_FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), HEADING_FONT_CN)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL_DARK)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(rpr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_heading_text(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"(`+|\*\*|__|~~|\*)", "", text)
    return html.unescape(text.strip())


def github_slug(text: str) -> str:
    value = clean_heading_text(text).lower()
    value = re.sub(r"[^\w\-\u4e00-\u9fff\s]", "", value)
    value = re.sub(r"\s+", "-", value).strip("-")
    return value


def strip_source_toc_and_meta(lines: list[str]) -> tuple[list[str], list[str], str]:
    title = PRODUCT
    first_h1 = None
    for idx, line in enumerate(lines):
        m = re.match(r"^#\s+(.+?)\s*$", line)
        if m:
            first_h1 = idx
            title = clean_heading_text(m.group(1))
            break

    meta: list[str] = []
    body = list(lines)
    if first_h1 is not None:
        del body[first_h1]
        pos = first_h1
        while pos < len(body) and not body[pos].strip():
            del body[pos]
        while pos < len(body):
            line = body[pos]
            if line.lstrip().startswith(">"):
                meta.append(re.sub(r"^\s*>\s?", "", line).rstrip())
                del body[pos]
                continue
            if not line.strip() and meta:
                del body[pos]
                continue
            break

    filtered: list[str] = []
    i = 0
    while i < len(body):
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", body[i])
        if m and clean_heading_text(m.group(2)).replace(" ", "") == "目录":
            toc_level = len(m.group(1))
            i += 1
            while i < len(body):
                if re.match(r"^\s*(---+|___+|\*\*\*+)\s*$", body[i]):
                    i += 1
                    break
                hm = re.match(r"^(#{1,6})\s+(.+?)\s*$", body[i])
                if hm and len(hm.group(1)) <= toc_level:
                    break
                i += 1
            continue
        filtered.append(body[i])
        i += 1
    return filtered, meta, title


def collect_headings(lines: list[str]) -> list[HeadingInfo]:
    headings: list[HeadingInfo] = []
    slug_counts: defaultdict[str, int] = defaultdict(int)
    active_fence: str | None = None
    for line in lines:
        fence_match = re.match(r"^\s*(```|~~~)([A-Za-z0-9_+.-]*)\s*$", line)
        if active_fence is not None:
            if line.strip() == active_fence:
                active_fence = None
            continue
        if fence_match:
            active_fence = fence_match.group(1)
            continue
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not m:
            continue
        source_level = len(m.group(1))
        text = clean_heading_text(m.group(2))
        if not text or text.replace(" ", "") == "目录":
            continue
        word_level = 1 if source_level <= 2 else min(5, source_level - 1)
        raw_slug = github_slug(text) or f"section-{len(headings) + 1}"
        slug_counts[raw_slug] += 1
        source_slug = raw_slug if slug_counts[raw_slug] == 1 else f"{raw_slug}-{slug_counts[raw_slug] - 1}"
        digest = hashlib.sha1(f"{len(headings)}:{text}".encode("utf-8")).hexdigest()[:8]
        bookmark = f"AXH_{len(headings) + 1}_{digest}"
        headings.append(HeadingInfo(source_level, word_level, text, bookmark, source_slug))
    return headings


def add_toc_and_guide(
    doc: Document,
    code: str,
    info: dict[str, object],
    meta_lines: list[str],
    headings: list[HeadingInfo],
) -> None:
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, cover=False)
    set_page_number_start(body_section, 1)
    configure_header_footer(body_section, code, str(info["title"]))

    top = doc.add_paragraph()
    add_bookmark(top, "AXIOM_TOP", 1)
    top.paragraph_format.space_after = Pt(0)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    kr = kicker.add_run(f"DOCUMENT {code}  /  REVIEWER GUIDE")
    set_run_font(kr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=9.0, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    tr = title.add_run("文档导读")
    set_run_font(tr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=22, color=NAVY, bold=True)

    purpose = doc.add_paragraph()
    purpose.style = doc.styles["AXIOM Quote"]
    set_paragraph_shading(purpose, CALLOUT_BG)
    set_paragraph_border(purpose, side="left", color=TEAL, size=22, space=6)
    pr = purpose.add_run("本册职责　")
    set_run_font(pr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=10.5, color=TEAL_DARK, bold=True)
    rr = purpose.add_run(str(info["purpose"]))
    set_run_font(rr, latin=BODY_FONT_LATIN, east_asia=BODY_FONT_CN, size=10.5, color=INK)

    if meta_lines:
        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(8)
        label.paragraph_format.space_after = Pt(5)
        lr = label.add_run("文档信息")
        set_run_font(lr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=11.0, color=NAVY, bold=True)
        for raw in meta_lines:
            if not raw.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Mm(3)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_inline(p, raw.rstrip().removesuffix("  "), SOURCE_DIR / f"{code}.md", None, None)

    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_before = Pt(12)
    toc_title.paragraph_format.space_after = Pt(4)
    add_bookmark(toc_title, "AXIOM_TOC", 2)
    ttr = toc_title.add_run("目录")
    set_run_font(ttr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=16.5, color=NAVY, bold=True)
    hint = doc.add_paragraph()
    hint.paragraph_format.space_after = Pt(7)
    hr = hint.add_run("Word 中可点击条目直接跳转；目录层级与正文书签保持一致。")
    set_run_font(hr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.5, color=MUTED)

    depth = int(info.get("toc_depth", 1))
    entries = [h for h in headings if h.word_level <= depth]
    major_index = 0
    for heading in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(0 if heading.word_level == 1 else 7)
        p.paragraph_format.first_line_indent = Mm(0)
        p.paragraph_format.space_before = Pt(1.5 if heading.word_level == 1 else 0)
        p.paragraph_format.space_after = Pt(3.5 if heading.word_level == 1 else 2)
        p.paragraph_format.keep_together = True
        if heading.word_level == 1:
            major_index += 1
        prefix = f"{major_index:02d}" if heading.word_level == 1 else "·"
        r = p.add_run(f"{prefix}  ")
        set_run_font(r, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=9.0, color=GOLD if heading.word_level == 1 else MUTED, bold=True)
        add_internal_hyperlink(p, heading.text, heading.bookmark, color=NAVY if heading.word_level == 1 else TEAL_DARK, bold=heading.word_level == 1)
        if heading.word_level == 1:
            set_paragraph_border(p, side="bottom", color=LINE, size=3, space=2)

    doc.add_page_break()


def ensure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    next_abs = max(existing_abs or [0]) + 1

    def create_abstract(abstract_id: int, kind: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(6):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
            lvl.append(num_fmt)
            lvl_text = OxmlElement("w:lvlText")
            if kind == "bullet":
                lvl_text.set(qn("w:val"), "•" if level % 3 == 0 else ("◦" if level % 3 == 1 else "▪"))
            else:
                lvl_text.set(qn("w:val"), f"%{level + 1}.")
            lvl.append(lvl_text)
            suff = OxmlElement("w:suff")
            suff.set(qn("w:val"), "space")
            lvl.append(suff)
            ppr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            left = 540 + level * 360
            tab.set(qn("w:pos"), str(left))
            tabs.append(tab)
            ppr.append(tabs)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(left))
            ind.set(qn("w:hanging"), "270")
            ppr.append(ind)
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:after"), "80")
            spacing.set(qn("w:line"), "300")
            spacing.set(qn("w:lineRule"), "auto")
            ppr.append(spacing)
            lvl.append(ppr)
            if kind == "bullet":
                rpr = OxmlElement("w:rPr")
                rf = OxmlElement("w:rFonts")
                rf.set(qn("w:ascii"), "Arial")
                rf.set(qn("w:hAnsi"), "Arial")
                rpr.append(rf)
                lvl.append(rpr)
            abstract.append(lvl)
        numbering.append(abstract)

    create_abstract(next_abs, "bullet")
    create_abstract(next_abs + 1, "decimal")
    return next_abs, next_abs + 1


def new_num_instance(doc: Document, abstract_id: int, start: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    if start != 1:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start))
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    numid = num_pr.find(qn("w:numId"))
    if numid is None:
        numid = OxmlElement("w:numId")
        num_pr.append(numid)
    numid.set(qn("w:val"), str(num_id))


def resolve_link_target(raw_target: str, source_path: Path) -> tuple[str, bool]:
    target = html.unescape(raw_target.strip().strip("<>"))
    target = re.sub(r'\s+["\'][^"\']*["\']\s*$', "", target)
    if target.startswith("#"):
        return target[1:], True
    if re.match(r"^(?:https?|mailto):", target, flags=re.I):
        return target, False
    path_part, sep, fragment = target.partition("#")
    local_path = (source_path.parent / path_part).resolve()
    if local_path.suffix.lower() == ".md" and local_path.parent == SOURCE_DIR.resolve():
        rewritten = local_path.with_suffix(".docx").name
    else:
        rewritten = os.path.relpath(local_path, OUTPUT_DIR).replace(os.sep, "/")
    if sep and fragment:
        rewritten += f"#{fragment}"
    return rewritten, False


INLINE_TOKEN = re.compile(
    r"(!\[[^\]]*\]\([^)]*\)|\[[^\]]+\]\([^)]*\)|`[^`]+`|\*\*.+?\*\*|~~.+?~~|(?<!\*)\*[^*\n]+?\*(?!\*))"
)
RAW_URL = re.compile(r"https?://[^\s<>()]+")


def add_plain_with_urls(paragraph, text: str, *, bold: bool, italic: bool, size: float | None = None, color: str | None = None) -> None:
    pos = 0
    for match in RAW_URL.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        add_external_hyperlink(paragraph, match.group(0), match.group(0), bold=bold, italic=italic)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_inline(
    paragraph,
    text: str,
    source_path: Path,
    slug_map: dict[str, str] | None,
    stats: BuildStats | None,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float | None = None,
    color: str | None = None,
) -> None:
    text = html.unescape(text).replace("<br/>", "\n").replace("<br />", "\n").replace("<br>", "\n")
    pos = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > pos:
            chunk = text[pos : match.start()]
            pieces = chunk.split("\n")
            for idx, piece in enumerate(pieces):
                if piece:
                    add_plain_with_urls(paragraph, piece.replace("\\|", "|"), bold=bold, italic=italic, size=size, color=color)
                if idx < len(pieces) - 1:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
        token = match.group(0)
        if token.startswith("!["):
            alt = re.match(r"!\[([^\]]*)\]", token)
            run = paragraph.add_run(alt.group(1) if alt else "[图片]")
            set_run_font(run, size=size, color=MUTED, italic=True)
        elif token.startswith("["):
            lm = re.match(r"\[([^\]]+)\]\(([^)]*)\)", token)
            if lm:
                label, raw_target = lm.group(1), lm.group(2)
                target, internal = resolve_link_target(raw_target, source_path)
                if internal:
                    anchor = (slug_map or {}).get(target, (slug_map or {}).get(github_slug(target), ""))
                    if anchor:
                        add_internal_hyperlink(paragraph, label, anchor, bold=bold)
                    else:
                        run = paragraph.add_run(label)
                        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
                else:
                    add_external_hyperlink(paragraph, label, target, bold=bold, italic=italic)
                    if stats:
                        stats.hyperlinks += 1
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin=CODE_FONT, east_asia=HEADING_FONT_CN, size=size or 9.0, color=NAVY_2, bold=bold, italic=italic)
            set_run_shading(run, CODE_BG)
        elif token.startswith("**") or token.startswith("__"):
            add_inline(paragraph, token[2:-2], source_path, slug_map, stats, bold=True, italic=italic, size=size, color=color)
        elif token.startswith("~~"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=MUTED, bold=bold, italic=italic)
            run.font.strike = True
        elif token.startswith("*"):
            add_inline(paragraph, token[1:-1], source_path, slug_map, stats, bold=bold, italic=True, size=size, color=color)
        pos = match.end()
    if pos < len(text):
        tail = text[pos:]
        pieces = tail.split("\n")
        for idx, piece in enumerate(pieces):
            if piece:
                add_plain_with_urls(paragraph, piece.replace("\\|", "|"), bold=bold, italic=italic, size=size, color=color)
            if idx < len(pieces) - 1:
                paragraph.add_run().add_break(WD_BREAK.LINE)


def set_run_shading(run, fill: str) -> None:
    rpr = run._element.get_or_add_rPr()
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def split_md_table_row(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|") and not value.endswith("\\|"):
        value = value[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    in_code = False
    tick_count = 0
    i = 0
    while i < len(value):
        ch = value[i]
        if escaped:
            current.append(ch)
            escaped = False
        elif ch == "\\":
            escaped = True
            current.append(ch)
        elif ch == "`":
            run = 1
            while i + run < len(value) and value[i + run] == "`":
                run += 1
            current.extend("`" * run)
            if not in_code:
                in_code = True
                tick_count = run
            elif run == tick_count:
                in_code = False
                tick_count = 0
            i += run - 1
        elif ch == "|" and not in_code:
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(ch)
        i += 1
    cells.append("".join(current).strip())
    return cells


def is_table_separator(line: str) -> bool:
    cells = split_md_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def visual_len(text: str) -> float:
    clean = re.sub(r"[`*_\[\]()<>]", "", text)
    return sum(1.0 if ord(ch) > 255 else 0.55 for ch in clean)


def table_widths(rows: list[list[str]], ncols: int) -> list[int]:
    if ncols <= 0:
        return []
    lengths = []
    for col in range(ncols):
        vals = [visual_len(row[col]) for row in rows if col < len(row)]
        sample = sorted(vals, reverse=True)[: min(8, len(vals))]
        score = (sum(sample) / max(1, len(sample))) + (max(vals or [1]) ** 0.45)
        lengths.append(max(3.0, score))
    if ncols == 2 and lengths[0] < lengths[1] * 0.55:
        raw = [0.30, 0.70]
    elif ncols == 3 and lengths[0] < max(lengths[1:]) * 0.6:
        raw = [0.22, 0.39, 0.39]
    else:
        roots = [math.sqrt(v) for v in lengths]
        total = sum(roots)
        raw = [v / total for v in roots]
    min_width = 520 if ncols >= 9 else (650 if ncols >= 7 else 820)
    widths = [max(min_width, round(TABLE_WIDTH_DXA * ratio)) for ratio in raw]
    while sum(widths) > TABLE_WIDTH_DXA:
        idx = max(range(ncols), key=lambda k: widths[k] - min_width)
        if widths[idx] <= min_width:
            break
        widths[idx] -= min(20, sum(widths) - TABLE_WIDTH_DXA)
    while sum(widths) < TABLE_WIDTH_DXA:
        idx = max(range(ncols), key=lambda k: lengths[k])
        widths[idx] += TABLE_WIDTH_DXA - sum(widths)
    if sum(widths) != TABLE_WIDTH_DXA:
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(
    doc: Document,
    rows: list[list[str]],
    source_path: Path,
    slug_map: dict[str, str],
    stats: BuildStats,
) -> None:
    if not rows:
        return
    ncols = max(len(row) for row in rows)
    normalized = []
    for row in rows:
        if len(row) > ncols:
            row = row[: ncols - 1] + [" | ".join(row[ncols - 1 :])]
        normalized.append(row + [""] * (ncols - len(row)))
    rows = normalized
    widths = table_widths(rows, ncols)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    font_size = 8.4 if ncols <= 4 else (7.6 if ncols <= 6 else 6.9)

    def fill_row(cells, values: list[str], row_index: int) -> None:
        for col, cell in enumerate(cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(
                p,
                values[col].replace("<br/>", "\n").replace("<br>", "\n"),
                source_path,
                slug_map,
                stats,
                bold=row_index == 0,
                size=font_size,
                color=WHITE if row_index == 0 else INK,
            )
            for run in p.runs:
                if row_index == 0:
                    run.font.color.rgb = rgb(WHITE)
                    run.bold = True
                if run.font.size is None or run.font.size.pt > font_size:
                    run.font.size = Pt(font_size)

    fill_row(table.rows[0].cells, rows[0], 0)
    repeat_table_header(table.rows[0])
    prevent_table_row_split(table.rows[0])
    for ridx, values in enumerate(rows[1:], start=1):
        row = table.add_row()
        fill_row(row.cells, values, ridx)
        prevent_table_row_split(row)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    stats.tables += 1
    stats.table_rows += len(rows)


def image_size_for_page(path: Path) -> tuple[float, float]:
    with Image.open(path) as im:
        width, height = im.size
    ratio = width / max(1, height)
    max_w, max_h = 6.35, 7.25
    if ratio >= max_w / max_h:
        w = max_w
        h = w / ratio
    else:
        h = max_h
        w = h * ratio
    # Avoid aggressively upscaling small screenshots.
    w = min(w, max(2.2, width / 110))
    h = w / ratio
    if h > max_h:
        h = max_h
        w = h * ratio
    return w, h


def set_inline_shape_alt(shape, description: str) -> None:
    inline = shape._inline
    doc_pr = inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_figure(doc: Document, path: Path, caption: str, stats: BuildStats) -> None:
    if not path.exists():
        p = doc.add_paragraph(style="AXIOM Quote")
        set_paragraph_shading(p, RISK_BG)
        set_paragraph_border(p, side="left", color=GOLD, size=20, space=6)
        r = p.add_run(f"图像资源未找到：{path}")
        set_run_font(r, size=9.5, color=INK)
        stats.missing_assets += 1
        return
    w, h = image_size_for_page(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(w), height=Inches(h))
    set_inline_shape_alt(shape, caption)
    stats.figures += 1
    cp = doc.add_paragraph(style="Caption")
    cr = cp.add_run(f"图 {stats.figures}　{caption}")
    set_run_font(cr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.5, color=MUTED)


def wrap_label(text: str, width: int = 17) -> list[str]:
    text = text.replace("<br/>", "\n").replace("<br>", "\n")
    lines: list[str] = []
    for part in text.splitlines() or [text]:
        if not part:
            lines.append("")
            continue
        line = ""
        units = 0.0
        for ch in part:
            step = 1.0 if ord(ch) > 255 else 0.55
            if line and units + step > width:
                lines.append(line)
                line = ch
                units = step
            else:
                line += ch
                units += step
        if line:
            lines.append(line)
    return lines[:5]


def parse_mermaid(source: str) -> tuple[str, dict[str, tuple[str, str]], list[tuple[str, str, str, str]], dict[str, list[str]]]:
    lines = [line.strip() for line in source.splitlines() if line.strip() and not line.strip().startswith("%%")]
    orientation = "TD"
    if lines:
        m = re.match(r"(?:flowchart|graph)\s+(LR|RL|TD|TB|BT)", lines[0], re.I)
        if m:
            orientation = m.group(1).upper()
    nodes: dict[str, tuple[str, str]] = {}
    edges: list[tuple[str, str, str, str]] = []
    groups: dict[str, list[str]] = defaultdict(list)
    active_group: str | None = None
    for raw in lines[1:]:
        sm = re.match(r"subgraph\s+\w+(?:\[\"([^\"]+)\"\]|\[([^\]]+)\])?", raw, re.I)
        if sm:
            active_group = sm.group(1) or sm.group(2) or "Group"
            continue
        if raw.lower() == "end":
            active_group = None
            continue
        for nm in re.finditer(r"\b([A-Za-z][\w]*)\s*(\[\"([^\"]+)\"\]|\[([^\]]+)\]|\{\"([^\"]+)\"\}|\{([^}]+)\})", raw):
            node_id = nm.group(1)
            label = next((g for g in nm.groups()[2:] if g is not None), node_id)
            shape = "decision" if "{" in nm.group(2) else "box"
            nodes[node_id] = (label, shape)
            if active_group and node_id not in groups[active_group]:
                groups[active_group].append(node_id)
        simplified = re.sub(
            r"\b([A-Za-z][\w]*)\s*(?:\[\"[^\"]+\"\]|\[[^\]]+\]|\{\"[^\"]+\"\}|\{[^}]+\})",
            r"\1",
            raw,
        )
        em = re.match(r"([A-Za-z][\w]*)\s*(?:--\s*\"([^\"]+)\"\s*-->|--\s*([^\-]+?)\s*-->|(-\.->|-->>|-->|---))\s*([A-Za-z][\w]*)", simplified)
        if em:
            src = em.group(1)
            label = (em.group(2) or em.group(3) or "").strip()
            style = em.group(4) or "-->"
            dst = em.group(5)
            edges.append((src, dst, label, style))
            nodes.setdefault(src, (src, "box"))
            nodes.setdefault(dst, (dst, "box"))
        elif "-->" in simplified or "-.->" in simplified or "-->>" in simplified:
            em2 = re.search(r"([A-Za-z][\w]*)\s*(-\.->|-->>|-->)\s*([A-Za-z][\w]*)", simplified)
            if em2:
                src, style, dst = em2.groups()
                edges.append((src, dst, "", style))
                nodes.setdefault(src, (src, "box"))
                nodes.setdefault(dst, (dst, "box"))
        if active_group:
            for node_id in re.findall(r"\b([A-Za-z][\w]*)\b", simplified):
                if node_id in nodes and node_id not in groups[active_group]:
                    groups[active_group].append(node_id)
    return orientation, nodes, edges, dict(groups)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[float, float], end: tuple[float, float], *, color=(55, 84, 109), dashed: bool = False) -> None:
    sx, sy = start
    ex, ey = end
    if dashed:
        parts = 12
        for i in range(parts):
            if i % 2 == 0:
                a = i / parts
                b = (i + 1) / parts
                draw.line((sx + (ex - sx) * a, sy + (ey - sy) * a, sx + (ex - sx) * b, sy + (ey - sy) * b), fill=color, width=5)
    else:
        draw.line((sx, sy, ex, ey), fill=color, width=5)
    angle = math.atan2(ey - sy, ex - sx)
    length = 18
    spread = 0.55
    p1 = (ex - length * math.cos(angle - spread), ey - length * math.sin(angle - spread))
    p2 = (ex - length * math.cos(angle + spread), ey - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill=color)


def wrap_text_pixels(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: float,
) -> list[str]:
    """Wrap mixed Chinese/Latin labels by rendered width, not character count."""
    text = text.replace("<br/>", "\n").replace("<br>", "\n")
    wrapped: list[str] = []
    for paragraph in text.splitlines() or [text]:
        if not paragraph:
            wrapped.append("")
            continue
        remaining = paragraph.strip()
        while remaining:
            if draw.textlength(remaining, font=font) <= max_width:
                wrapped.append(remaining)
                break
            cut = 1
            last_break = 0
            for idx in range(1, len(remaining) + 1):
                candidate = remaining[:idx]
                if draw.textlength(candidate, font=font) > max_width:
                    break
                cut = idx
                if remaining[idx - 1] in " /、，,；;·|":
                    last_break = idx
            if last_break >= max(2, cut // 2):
                cut = last_break
            wrapped.append(remaining[:cut].strip())
            remaining = remaining[cut:].strip()
    return wrapped or [""]


def fitted_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    font: ImageFont.FreeTypeFont,
    max_width: float,
    max_height: float,
    *,
    min_size: int = 18,
    spacing: int = 7,
) -> tuple[ImageFont.FreeTypeFont, list[str], tuple[int, int, int, int]]:
    """Return wrapped text and the largest font that fits the available box."""
    candidate = font
    while True:
        lines = wrap_text_pixels(draw, text, candidate, max_width)
        block = "\n".join(lines)
        bbox = draw.multiline_textbbox((0, 0), block, font=candidate, spacing=spacing, align="center")
        if bbox[2] - bbox[0] <= max_width + 1 and bbox[3] - bbox[1] <= max_height + 1:
            return candidate, lines, bbox
        size = int(getattr(candidate, "size", min_size)) - 2
        if size < min_size:
            return candidate, lines, bbox
        candidate = candidate.font_variant(size=size)


def draw_node(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], label: str, shape: str, font, *, fill=(255, 255, 255), outline=(20, 168, 154)) -> None:
    x1, y1, x2, y2 = box
    if shape == "decision":
        points = [((x1 + x2) // 2, y1), (x2, (y1 + y2) // 2), ((x1 + x2) // 2, y2), (x1, (y1 + y2) // 2)]
        draw.polygon(points, fill=(255, 249, 235), outline=(214, 166, 75))
        draw.line(points + [points[0]], fill=(214, 166, 75), width=4)
    else:
        draw.rounded_rectangle(box, radius=22, fill=fill, outline=outline, width=4)
    box_w, box_h = x2 - x1, y2 - y1
    max_text_w = box_w - 54 if shape == "box" else box_w * 0.61
    max_text_h = box_h - 28 if shape == "box" else box_h * 0.58
    fitted_font, lines, bbox = fitted_text(draw, label, font, max_text_w, max_text_h)
    block = "\n".join(lines)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - tw) / 2 - bbox[0], (y1 + y2 - th) / 2 - bbox[1]),
        block,
        font=fitted_font,
        fill=(23, 34, 53),
        spacing=7,
        align="center",
    )


def render_sequence_mermaid(source: str, out: Path) -> None:
    lines = [line.strip() for line in source.splitlines() if line.strip()]
    participants: list[tuple[str, str]] = []
    events: list[tuple[str, str, str, str]] = []
    notes: list[tuple[str, str]] = []
    for line in lines:
        pm = re.match(r"participant\s+(\w+)\s+as\s+(.+)", line, re.I)
        if pm:
            participants.append((pm.group(1), pm.group(2)))
            continue
        mm = re.match(r"(\w+)(-->>|->>|-->|->)(\w+)\s*:\s*(.+)", line)
        if mm:
            events.append((mm.group(1), mm.group(3), mm.group(4), mm.group(2)))
            continue
        nm = re.match(r"Note\s+over\s+(\w+)\s*:\s*(.+)", line, re.I)
        if nm:
            notes.append((nm.group(1), nm.group(2)))
            events.append((nm.group(1), nm.group(1), nm.group(2), "note"))
    if not participants:
        participants = [("A", "参与者 A"), ("B", "参与者 B")]
    width = max(1500, 460 * len(participants) + 220)
    height = 330 + 145 * max(1, len(events))
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(FONT_BOLD, 34)
    font = load_font(FONT_REGULAR, 30)
    small = load_font(FONT_REGULAR, 26)
    xs = {pid: 190 + i * ((width - 380) / max(1, len(participants) - 1)) for i, (pid, _) in enumerate(participants)}
    for pid, label in participants:
        x = xs[pid]
        actor_box = (x - 170, 42, x + 170, 145)
        draw.rounded_rectangle(actor_box, radius=18, fill=(16, 36, 62), outline=(16, 36, 62), width=3)
        actor_font, actor_lines, actor_bbox = fitted_text(draw, label, title_font, 300, 76, min_size=22, spacing=4)
        actor_text = "\n".join(actor_lines)
        actor_w, actor_h = actor_bbox[2] - actor_bbox[0], actor_bbox[3] - actor_bbox[1]
        draw.multiline_text(
            (x - actor_w / 2 - actor_bbox[0], 93.5 - actor_h / 2 - actor_bbox[1]),
            actor_text,
            font=actor_font,
            fill="white",
            spacing=4,
            align="center",
        )
        draw.line((x, 145, x, height - 70), fill=(170, 183, 196), width=3)
    y = 220
    for src, dst, label, kind in events:
        if kind == "note":
            x = xs.get(src, width / 2)
            box_x1 = max(35, min(width - 495, x - 230))
            box_x2 = box_x1 + 460
            lines2 = wrap_text_pixels(draw, label, small, 410)
            bbox2 = draw.multiline_textbbox((0, 0), "\n".join(lines2), font=small, spacing=6, align="center")
            box_h = max(92, bbox2[3] - bbox2[1] + 42)
            draw.rounded_rectangle((box_x1, y - 20, box_x2, y - 20 + box_h), radius=12, fill=(255, 248, 230), outline=(214, 166, 75), width=3)
            draw.multiline_text(
                ((box_x1 + box_x2 - (bbox2[2] - bbox2[0])) / 2 - bbox2[0], y + (box_h - (bbox2[3] - bbox2[1])) / 2 - 20 - bbox2[1]),
                "\n".join(lines2),
                font=small,
                fill=(76, 63, 35),
                spacing=6,
                align="center",
            )
            y += box_h + 28
            continue
        x1, x2 = xs.get(src, 150), xs.get(dst, width - 150)
        start = (x1, y)
        end = (x2 - (12 if x2 > x1 else -12), y)
        draw_arrow(draw, start, end, dashed="--" in kind)
        message_w = max(120, abs(x2 - x1) - 70)
        msg_font, msg_lines, msg_bbox = fitted_text(draw, label, font, message_w, 64, min_size=20, spacing=4)
        msg_text = "\n".join(msg_lines)
        tw, th = msg_bbox[2] - msg_bbox[0], msg_bbox[3] - msg_bbox[1]
        mx = (x1 + x2) / 2
        draw.rounded_rectangle((mx - tw / 2 - 10, y - th - 22, mx + tw / 2 + 10, y - 4), radius=5, fill="white")
        draw.multiline_text((mx - tw / 2 - msg_bbox[0], y - th - 17 - msg_bbox[1]), msg_text, font=msg_font, fill=(23, 34, 53), spacing=4, align="center")
        y += 115
    image.save(out)


def render_grouped_mermaid(
    nodes: dict[str, tuple[str, str]],
    edges: list[tuple[str, str, str, str]],
    groups: dict[str, list[str]],
    out: Path,
) -> None:
    group_items = list(groups.items())
    width = 1900
    band_heights = []
    for _, ids in group_items:
        rows = math.ceil(max(1, len(ids)) / 4)
        band_heights.append(125 + 122 * rows + 32 * max(0, rows - 1))
    height = 100 + sum(band_heights) + 55 * max(0, len(group_items) - 1)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    group_font = load_font(FONT_BOLD, 32)
    node_font = load_font(FONT_REGULAR, 27)
    positions: dict[str, tuple[int, int, int, int]] = {}
    y = 50
    fills = [(238, 248, 246), (237, 243, 249), (247, 248, 250), (255, 248, 233), (242, 246, 250)]
    pending_nodes: list[tuple[str, tuple[int, int, int, int]]] = []
    for gi, (group, ids) in enumerate(group_items):
        band_h = band_heights[gi]
        draw.rounded_rectangle((45, y, width - 45, y + band_h), radius=22, fill=fills[gi % len(fills)], outline=(185, 199, 211), width=3)
        draw.text((75, y + 20), group, font=group_font, fill=(16, 36, 62))
        cols = min(4, max(1, len(ids)))
        cell_w = (width - 190) / cols
        for idx, node_id in enumerate(ids):
            row, col = divmod(idx, cols)
            x1 = 85 + int(col * cell_w)
            y1 = y + 76 + row * 154
            x2 = x1 + int(cell_w - 34)
            y2 = y1 + 118
            positions[node_id] = (x1, y1, x2, y2)
            pending_nodes.append((node_id, positions[node_id]))
        y += band_h + 55
    for src, dst, label, style in edges:
        if src not in positions or dst not in positions:
            continue
        a, b = positions[src], positions[dst]
        start = ((a[0] + a[2]) / 2, a[3])
        end = ((b[0] + b[2]) / 2, b[1])
        if end[1] <= start[1]:
            continue
        draw_arrow(draw, start, end, color=(96, 119, 139), dashed="-." in style)
    for node_id, box in pending_nodes:
        label, shape = nodes.get(node_id, (node_id, "box"))
        draw_node(draw, box, label, shape, node_font, fill=(255, 255, 255), outline=(20, 168, 154))
    image.save(out)


def render_flow_mermaid(source: str, out: Path) -> None:
    orientation, nodes, edges, groups = parse_mermaid(source)
    if groups:
        render_grouped_mermaid(nodes, edges, groups, out)
        return
    ids = list(nodes)
    if not ids:
        image = Image.new("RGB", (1400, 320), "white")
        draw = ImageDraw.Draw(image)
        draw.text((60, 80), source[:180], font=load_font(FONT_REGULAR, 32), fill=(23, 34, 53))
        image.save(out)
        return

    incoming: defaultdict[str, int] = defaultdict(int)
    outgoing: defaultdict[str, list[str]] = defaultdict(list)
    for src, dst, _, _ in edges:
        if dst not in outgoing[src]:
            outgoing[src].append(dst)
            incoming[dst] += 1
        incoming.setdefault(src, incoming.get(src, 0))
    roots = [node for node in ids if incoming[node] == 0] or [ids[0]]
    levels: dict[str, int] = {}
    queue = deque((root, 0) for root in roots)
    while queue:
        node, level = queue.popleft()
        if node in levels:
            continue
        levels[node] = level
        for nxt in outgoing.get(node, []):
            if nxt not in levels:
                queue.append((nxt, level + 1))
    for node in ids:
        if node not in levels:
            neighbors = [levels[src] + 1 for src, dst, _, _ in edges if dst == node and src in levels]
            levels[node] = min(neighbors) if neighbors else max(levels.values(), default=-1) + 1
    grouped: defaultdict[int, list[str]] = defaultdict(list)
    for node in ids:
        grouped[levels[node]].append(node)
    layer_keys = sorted(grouped)
    horizontal = orientation in {"LR", "RL"}
    # Word pages are portrait.  Long left-to-right pipelines become unreadably
    # thin when forced into one row, so print them top-to-bottom after four
    # stages while preserving the exact graph topology and labels.
    if horizontal and len(layer_keys) > 4:
        horizontal = False
    node_w, node_h = 390, 152
    layer_gap, row_gap = 92, 64
    max_per_layer = max(len(grouped[k]) for k in layer_keys)
    if horizontal:
        width = 120 + len(layer_keys) * node_w + max(0, len(layer_keys) - 1) * layer_gap
        height = 120 + max_per_layer * node_h + max(0, max_per_layer - 1) * row_gap
    else:
        width = 120 + max_per_layer * node_w + max(0, max_per_layer - 1) * row_gap
        height = 120 + len(layer_keys) * node_h + max(0, len(layer_keys) - 1) * layer_gap
    width = max(width, 1200)
    height = max(height, 560)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = load_font(FONT_REGULAR, 28)
    edge_font = load_font(FONT_REGULAR, 22)
    positions: dict[str, tuple[int, int, int, int]] = {}
    for li, key in enumerate(layer_keys):
        layer = grouped[key]
        for ri, node in enumerate(layer):
            if horizontal:
                x = 60 + li * (node_w + layer_gap)
                total_h = len(layer) * node_h + (len(layer) - 1) * row_gap
                y = (height - total_h) // 2 + ri * (node_h + row_gap)
            else:
                total_w = len(layer) * node_w + (len(layer) - 1) * row_gap
                x = (width - total_w) // 2 + ri * (node_w + row_gap)
                y = 60 + li * (node_h + layer_gap)
            positions[node] = (x, y, x + node_w, y + node_h)
    for src, dst, label, style in edges:
        if src not in positions or dst not in positions:
            continue
        a, b = positions[src], positions[dst]
        if horizontal:
            if b[0] >= a[2]:
                start, end = (a[2], (a[1] + a[3]) / 2), (b[0], (b[1] + b[3]) / 2)
            else:
                start, end = (a[0], (a[1] + a[3]) / 2), (b[2], (b[1] + b[3]) / 2)
        else:
            if b[1] >= a[3]:
                start, end = ((a[0] + a[2]) / 2, a[3]), ((b[0] + b[2]) / 2, b[1])
            else:
                start, end = ((a[0] + a[2]) / 2, a[1]), ((b[0] + b[2]) / 2, b[3])
        draw_arrow(draw, start, end, dashed="-." in style)
        if label:
            mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
            bb = draw.textbbox((0, 0), label, font=edge_font)
            tw, th = bb[2] - bb[0], bb[3] - bb[1]
            draw.rounded_rectangle((mx - tw / 2 - 8, my - th / 2 - 5, mx + tw / 2 + 8, my + th / 2 + 5), radius=6, fill="white")
            draw.text((mx - tw / 2, my - th / 2), label, font=edge_font, fill=(76, 91, 105))
    for node, box in positions.items():
        label, shape = nodes[node]
        draw_node(draw, box, label, shape, font)
    image.save(out)


def render_mermaid(source: str) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256(source.encode("utf-8")).hexdigest()[:16]
    out = DIAGRAM_DIR / f"mermaid-{digest}.png"
    if out.exists():
        return out
    if source.lstrip().startswith("sequenceDiagram"):
        render_sequence_mermaid(source, out)
    else:
        render_flow_mermaid(source, out)
    return out


def add_code_block(doc: Document, code: str, language: str, stats: BuildStats) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after = Pt(2)
    name = "结构示意" if language in {"", "text", "plaintext"} else f"代码示例 · {language}"
    lr = label.add_run(name.upper() if language and language.isascii() else name)
    set_run_font(lr, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=7.8, color=TEAL_DARK, bold=True)
    lines = code.splitlines() or [""]
    for idx, line in enumerate(lines):
        p = doc.add_paragraph(style="AXIOM Code")
        set_paragraph_shading(p, CODE_BG)
        set_paragraph_border(p, side="left", color=TEAL if idx == 0 else LINE, size=16 if idx == 0 else 4, space=4)
        if idx == 0:
            p.paragraph_format.space_before = Pt(1)
        if idx == len(lines) - 1:
            p.paragraph_format.space_after = Pt(6)
        r = p.add_run(line if line else " ")
        set_run_font(r, latin=CODE_FONT, east_asia=HEADING_FONT_CN, size=8.0, color=INK)
    stats.code_blocks += 1


def add_body_paragraph(doc: Document, text: str, source_path: Path, slug_map: dict[str, str], stats: BuildStats) -> None:
    p = doc.add_paragraph(style="AXIOM Body")
    add_inline(p, text, source_path, slug_map, stats)
    stats.paragraphs += 1


def add_quote(doc: Document, lines: list[str], source_path: Path, slug_map: dict[str, str], stats: BuildStats) -> None:
    text = "\n".join(re.sub(r"^\s*>\s?", "", line).rstrip().removesuffix("  ") for line in lines)
    p = doc.add_paragraph(style="AXIOM Quote")
    set_paragraph_shading(p, CALLOUT_BG)
    set_paragraph_border(p, side="left", color=TEAL, size=22, space=6)
    add_inline(p, text, source_path, slug_map, stats)
    stats.paragraphs += 1


def is_image_only(line: str) -> re.Match[str] | None:
    return re.fullmatch(r"\s*!\[([^\]]*)\]\(([^)]+)\)\s*", line)


def is_list_line(line: str) -> re.Match[str] | None:
    return re.match(r"^(\s*)([-+*]|\d+[.)])\s+(.+)$", line)


def is_structural(lines: list[str], index: int) -> bool:
    line = lines[index]
    if not line.strip():
        return True
    if re.match(r"^#{1,6}\s+", line):
        return True
    if re.match(r"^\s*(?:```|~~~)", line):
        return True
    if re.match(r"^\s*(?:---+|___+|\*\*\*+)\s*$", line):
        return True
    if line.lstrip().startswith(">"):
        return True
    if is_list_line(line):
        return True
    if is_image_only(line):
        return True
    if line.strip().startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
        return True
    if re.fullmatch(r"\s*</?[A-Za-z][^>]*>\s*", line):
        return True
    return False


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_border(p, side="bottom", color=LINE, size=6, space=2)


def convert_body(
    doc: Document,
    source_path: Path,
    lines: list[str],
    headings: list[HeadingInfo],
    stats: BuildStats,
) -> None:
    slug_map = {heading.source_slug: heading.bookmark for heading in headings}
    for heading in headings:
        slug_map.setdefault(github_slug(heading.text), heading.bookmark)
    heading_index = 0
    bullet_abstract, number_abstract = ensure_numbering(doc)
    active_kind: str | None = None
    active_num_id: int | None = None
    last_heading = "流程与结构示意"
    first_heading = True
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            active_kind = None
            active_num_id = None
            i += 1
            continue

        hm = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if hm:
            if heading_index >= len(headings):
                i += 1
                continue
            heading = headings[heading_index]
            heading_index += 1
            p = doc.add_paragraph(style=f"Heading {heading.word_level}")
            if first_heading and heading.word_level == 1:
                p.paragraph_format.page_break_before = False
                first_heading = False
            elif heading.word_level == 1 and re.match(
                r"^(?:附录|第[一二三四五六七八九十0-9]+(?:篇|部分))", heading.text
            ):
                p.paragraph_format.page_break_before = True
            add_inline(p, heading.text, source_path, slug_map, stats, bold=True)
            add_bookmark(p, heading.bookmark, 100 + heading_index)
            if heading.word_level == 1:
                set_paragraph_border(p, side="bottom", color=GOLD, size=12, space=5)
            stats.headings += 1
            last_heading = heading.text
            active_kind = None
            active_num_id = None
            i += 1
            continue

        fm = re.match(r"^\s*(```|~~~)([A-Za-z0-9_+.-]*)\s*$", line)
        if fm:
            fence, language = fm.group(1), fm.group(2).lower()
            end = i + 1
            while end < len(lines) and lines[end].strip() != fence:
                end += 1
            if end >= len(lines):
                add_body_paragraph(doc, line, source_path, slug_map, stats)
                i += 1
                continue
            content = "\n".join(lines[i + 1 : end])
            if language == "mermaid":
                diagram = render_mermaid(content)
                add_figure(doc, diagram, f"{last_heading}（流程/架构图）", stats)
                stats.mermaid += 1
            else:
                add_code_block(doc, content, language, stats)
            active_kind = None
            active_num_id = None
            i = end + 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = split_md_table_row(line)
            rows = [header]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_md_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows, source_path, slug_map, stats)
            active_kind = None
            active_num_id = None
            continue

        image_match = is_image_only(line)
        if image_match:
            alt, target = image_match.group(1), image_match.group(2).strip().strip("<>")
            path = (source_path.parent / target).resolve()
            add_figure(doc, path, alt or last_heading, stats)
            active_kind = None
            active_num_id = None
            i += 1
            continue

        if line.lstrip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                quote_lines.append(lines[i])
                i += 1
            add_quote(doc, quote_lines, source_path, slug_map, stats)
            active_kind = None
            active_num_id = None
            continue

        lm = is_list_line(line)
        if lm:
            indent, marker, content = lm.groups()
            level = min(5, max(0, len(indent.replace("\t", "    ")) // 2))
            checklist = re.match(r"^\[([ xX])\]\s*(.*)$", content)
            if checklist:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
                p.paragraph_format.first_line_indent = Inches(-0.19)
                p.paragraph_format.space_after = Pt(4)
                icon = "☑" if checklist.group(1).lower() == "x" else "☐"
                ir = p.add_run(f"{icon}  ")
                set_run_font(ir, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=10.0, color=TEAL_DARK, bold=True)
                add_inline(p, checklist.group(2), source_path, slug_map, stats)
                active_kind = "check"
                active_num_id = None
            else:
                kind = "number" if marker[0].isdigit() else "bullet"
                start = int(re.match(r"\d+", marker).group(0)) if kind == "number" else 1
                if active_kind != kind or active_num_id is None:
                    active_num_id = new_num_instance(doc, number_abstract if kind == "number" else bullet_abstract, start=start)
                active_kind = kind
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.25
                apply_numbering(p, active_num_id, level)
                add_inline(p, content, source_path, slug_map, stats)
            stats.paragraphs += 1
            i += 1
            continue

        if re.match(r"^\s*(?:---+|___+|\*\*\*+)\s*$", line):
            add_horizontal_rule(doc)
            active_kind = None
            active_num_id = None
            i += 1
            continue

        if re.fullmatch(r"\s*</?[A-Za-z][^>]*>\s*", line):
            add_code_block(doc, line.strip(), "xml", stats)
            active_kind = None
            active_num_id = None
            i += 1
            continue

        paragraph_lines = [line.rstrip()]
        i += 1
        while i < len(lines) and not is_structural(lines, i):
            paragraph_lines.append(lines[i].rstrip())
            i += 1
        joined = ""
        for part in paragraph_lines:
            hard_break = part.endswith("  ")
            clean = part[:-2] if hard_break else part
            if joined and not joined.endswith("\n"):
                joined += " "
            joined += clean.strip()
            if hard_break:
                joined += "\n"
        add_body_paragraph(doc, joined, source_path, slug_map, stats)
        active_kind = None
        active_num_id = None

    bottom = doc.add_paragraph()
    add_bookmark(bottom, "AXIOM_BOTTOM", 999999)
    bottom.paragraph_format.space_before = Pt(12)
    bottom.paragraph_format.space_after = Pt(0)
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bottom.add_run("— 文档结束 —")
    set_run_font(br, latin=HEADING_FONT_LATIN, east_asia=HEADING_FONT_CN, size=8.0, color=MUTED)


def scrub_metadata(path: Path) -> None:
    temp = path.with_name(path.stem + ".scrubbed.docx")
    skill_script = Path(
        r"C:\Users\why\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents\scripts\privacy_scrub.py"
    )
    if skill_script.exists():
        result = subprocess.run(
            [sys.executable, str(skill_script), str(path), "--out", str(temp)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        if result.returncode == 0 and temp.exists():
            # Windows preview/indexing processes can deny an atomic replace even
            # when a normal overwrite is available.  Copying the scrubbed bytes
            # keeps the operation deterministic without depending on rename
            # semantics, then removes only this builder-owned temporary file.
            shutil.copy2(temp, path)
            temp.unlink(missing_ok=True)
            return
    # Fallback: use neutral project-level core properties.
    doc = Document(str(path))
    doc.core_properties.author = TEAM_NAME
    doc.core_properties.last_modified_by = TEAM_NAME
    doc.core_properties.comments = ""
    doc.save(str(path))


def audit_docx(path: Path, stats: BuildStats) -> dict[str, object]:
    issues: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8")
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
        numbering_xml = zf.read("word/numbering.xml").decode("utf-8") if "word/numbering.xml" in names else ""
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8") if "word/_rels/document.xml.rels" in names else ""
        media = [name for name in names if name.startswith("word/media/")]
        if "[[TOC]]" in document_xml or "```" in document_xml or "~~~" in document_xml:
            issues.append("Markdown/TOC placeholder leaked into document XML")
        if "锛?" in document_xml or "鏂囨" in document_xml:
            issues.append("possible mojibake detected")
        if "w:tblLayout w:type=\"fixed\"" not in document_xml and stats.tables:
            issues.append("fixed table layout missing")
        if "w:abstractNum" not in numbering_xml:
            issues.append("real numbering definitions missing")
        if BODY_FONT_CN not in styles_xml or HEADING_FONT_CN not in styles_xml:
            issues.append("Chinese style fonts missing")
        if len(media) < stats.figures + 1:
            issues.append(f"embedded media count {len(media)} below expected {stats.figures + 1}")
        hyperlink_count = rels_xml.count("relationships/hyperlink")

    doc = Document(str(path))
    if len(doc.sections) < 2:
        issues.append("cover/body section split missing")
    for section in doc.sections:
        if abs(section.page_width.mm - PAGE_W_MM) > 0.6 or abs(section.page_height.mm - PAGE_H_MM) > 0.6:
            issues.append("non-A4 section detected")
            break
    for table_index, table in enumerate(doc.tables):
        if not table.rows or not table.columns:
            issues.append(f"empty table at index {table_index}")
            break
    return {
        "file": path.name,
        "bytes": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "embedded_media": len(media),
        "hyperlink_relationships": hyperlink_count,
        "issues": issues,
    }


def build_one(source_path: Path) -> tuple[Path, BuildStats, dict[str, object]]:
    code = source_path.stem[:2]
    if code not in DOC_INFO:
        raise ValueError(f"No document metadata for {source_path.name}")
    info = DOC_INFO[code]
    raw = source_path.read_text(encoding="utf-8-sig")
    raw_lines = raw.splitlines()
    lines, meta_lines, source_title = strip_source_toc_and_meta(raw_lines)
    headings = collect_headings(lines)
    stats = BuildStats(code=code, source_lines=len(raw_lines), source_chars=len(raw))

    doc = Document()
    set_doc_defaults(doc)
    configure_styles(doc)
    set_update_fields(doc)
    doc.core_properties.title = f"{PRODUCT} {str(info['title']).replace(chr(10), ' ')}"
    doc.core_properties.subject = COMPETITION
    doc.core_properties.keywords = f"{PRODUCT}, 软件杯, A3, {TEAM_NAME}, {code}"
    doc.core_properties.author = TEAM_NAME
    doc.core_properties.last_modified_by = TEAM_NAME

    # Remove the empty starter paragraph only after all styles/sections exist.
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    add_cover(doc, code, info)
    add_toc_and_guide(doc, code, info, meta_lines, headings)
    convert_body(doc, source_path, lines, headings, stats)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{source_path.stem}.docx"
    raw_path = WORK_DIR / f"{source_path.stem}.raw.docx"
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(raw_path))
    shutil.copy2(raw_path, output_path)
    scrub_metadata(output_path)
    stats.output_bytes = output_path.stat().st_size
    audit = audit_docx(output_path, stats)
    return output_path, stats, audit


def discover_sources(codes: set[str] | None = None) -> list[Path]:
    sources = sorted(SOURCE_DIR.glob("[0-1][0-9]-*.md"))
    sources = [path for path in sources if path.stem[:2] in DOC_INFO]
    if codes:
        sources = [path for path in sources if path.stem[:2] in codes]
    return sources


def inventory(sources: Iterable[Path]) -> list[dict[str, object]]:
    result = []
    for path in sources:
        raw = path.read_text(encoding="utf-8-sig")
        lines, meta, _ = strip_source_toc_and_meta(raw.splitlines())
        headings = collect_headings(lines)
        mermaid = len(re.findall(r"(?m)^\s*(?:```|~~~)mermaid\s*$", raw))
        images = len(re.findall(r"!\[[^\]]*\]\([^)]*\)", raw))
        tables = 0
        max_cols = 0
        split = raw.splitlines()
        for i in range(len(split) - 1):
            if split[i].strip().startswith("|") and is_table_separator(split[i + 1]):
                tables += 1
                max_cols = max(max_cols, len(split_md_table_row(split[i])))
        result.append(
            {
                "file": path.name,
                "lines": len(raw.splitlines()),
                "chars": len(raw),
                "meta_lines": len(meta),
                "headings": len(headings),
                "toc_entries": len([h for h in headings if h.word_level <= int(DOC_INFO[path.stem[:2]]["toc_depth"])]),
                "tables": tables,
                "max_table_columns": max_cols,
                "images": images,
                "mermaid": mermaid,
            }
        )
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--only", help="Comma-separated document codes, e.g. 00,09")
    parser.add_argument("--inventory", action="store_true", help="Print source inventory without building")
    args = parser.parse_args()
    codes = {part.strip().zfill(2) for part in args.only.split(",")} if args.only else None
    sources = discover_sources(codes)
    if not sources:
        raise SystemExit("No matching submission Markdown files found")
    if args.inventory:
        print(json.dumps(inventory(sources), ensure_ascii=False, indent=2))
        return 0

    all_stats = []
    all_audits = []
    for source in sources:
        print(f"[build] {source.name}", flush=True)
        output, stats, audit = build_one(source)
        all_stats.append(stats.__dict__)
        all_audits.append(audit)
        print(f"[done]  {output.name}  {output.stat().st_size / 1024 / 1024:.2f} MiB", flush=True)
        if audit["issues"]:
            print(f"[audit] {audit['issues']}", flush=True)
    report = {
        "preset": "compact_reference_guide",
        "overrides": ["A4_CN_SUBMISSION", "CN_TECH_TYPE", "AXIOM_BRAND", "DENSE_EVIDENCE_TABLE"],
        "competition": COMPETITION,
        "team": TEAM_NAME,
        "documents": all_stats,
        "audits": all_audits,
    }
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    (WORK_DIR / "build-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 1 if any(audit["issues"] for audit in all_audits) else 0


if __name__ == "__main__":
    raise SystemExit(main())
