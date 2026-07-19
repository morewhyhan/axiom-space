#!/usr/bin/env python3
"""Structural, fidelity, and rendered-page audit for the submission DOCX set."""

from __future__ import annotations

import html
import json
import os
import re
import sys
import zipfile
from collections import defaultdict
from pathlib import Path
from xml.etree import ElementTree as ET

import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.oxml.ns import qn

import build_submission_docx as builder


ROOT = builder.REPO_ROOT
SRC = builder.SOURCE_DIR
OUT = builder.OUTPUT_DIR
RENDERS = builder.WORK_DIR / "final-render"
RERENDERS = builder.WORK_DIR / "final-render-v2"
CONTACTS = builder.WORK_DIR / "contact-sheets"
REPORT_PATH = builder.WORK_DIR / "final-audit.json"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "dc": "http://purl.org/dc/elements/1.1/",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
}


def norm(value: str) -> str:
    value = html.unescape(value)
    value = value.replace("\\|", "|")
    return "".join(ch.lower() for ch in value if ch.isalnum() or ch == "_")


def visible_markdown(value: str) -> str:
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"(?<!_)__(?!_)(.+?)(?<!_)__(?!_)", r"\1", value)
    value = re.sub(r"(`+|\*\*|~~)", "", value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", value)
    value = value.replace("<br/>", " ").replace("<br />", " ").replace("<br>", " ")
    return value.strip()


def source_snippets(path: Path) -> list[str]:
    raw = path.read_text(encoding="utf-8-sig")
    lines, meta, _ = builder.strip_source_toc_and_meta(raw.splitlines())
    snippets: list[str] = []
    for item in meta:
        cleaned = visible_markdown(item.rstrip().removesuffix("  "))
        if len(norm(cleaned)) >= 6:
            snippets.append(cleaned)

    active_fence: str | None = None
    fence_language = ""
    for line in lines:
        fm = re.match(r"^\s*(```|~~~)([A-Za-z0-9_+.-]*)\s*$", line)
        if active_fence is not None:
            if line.strip() == active_fence:
                active_fence = None
                fence_language = ""
                continue
            if fence_language == "mermaid":
                continue
            cleaned = line.rstrip()
            if len(norm(cleaned)) >= 6:
                snippets.append(cleaned)
            continue
        if fm:
            active_fence = fm.group(1)
            fence_language = fm.group(2).lower()
            continue
        stripped = line.strip()
        if not stripped or re.fullmatch(r"(?:---+|___+|\*\*\*+)", stripped):
            continue
        if builder.is_table_separator(line):
            continue
        if stripped.startswith("|"):
            for cell in builder.split_md_table_row(stripped):
                cleaned = visible_markdown(cell)
                if len(norm(cleaned)) >= 6:
                    snippets.append(cleaned)
            continue
        image = builder.is_image_only(line)
        if image:
            if len(norm(image.group(1))) >= 6:
                snippets.append(image.group(1))
            continue
        cleaned = re.sub(r"^#{1,6}\s+", "", line)
        cleaned = re.sub(r"^\s*>\s?", "", cleaned)
        cleaned = re.sub(r"^\s*(?:[-+*]|\d+[.)])\s+", "", cleaned)
        cleaned = re.sub(r"^\[([ xX])\]\s*", "", cleaned)
        cleaned = visible_markdown(cleaned.rstrip().removesuffix("  "))
        if len(norm(cleaned)) >= 6:
            snippets.append(cleaned)
    return snippets


def document_visible_text(zf: zipfile.ZipFile) -> str:
    root = ET.fromstring(zf.read("word/document.xml"))
    texts = [node.text or "" for node in root.findall(".//w:t", NS)]
    return "".join(texts)


def relationship_targets(zf: zipfile.ZipFile) -> list[str]:
    path = "word/_rels/document.xml.rels"
    if path not in zf.namelist():
        return []
    root = ET.fromstring(zf.read(path))
    values = []
    for rel in root.findall("pr:Relationship", NS):
        rel_type = rel.attrib.get("Type", "")
        if rel_type.endswith("/hyperlink"):
            values.append(rel.attrib.get("Target", ""))
    return values


def table_geometry_issues(document_xml: bytes) -> list[str]:
    root = ET.fromstring(document_xml)
    issues = []
    for index, table in enumerate(root.findall(".//w:tbl", NS), start=1):
        layout = table.find("./w:tblPr/w:tblLayout", NS)
        if layout is None or layout.attrib.get(qn("w:type")) != "fixed":
            issues.append(f"table {index}: non-fixed layout")
        tbl_w = table.find("./w:tblPr/w:tblW", NS)
        if tbl_w is None or tbl_w.attrib.get(qn("w:type")) != "dxa":
            issues.append(f"table {index}: non-DXA table width")
        grid = table.findall("./w:tblGrid/w:gridCol", NS)
        grid_widths = [int(col.attrib.get(qn("w:w"), "0")) for col in grid]
        if not grid_widths or sum(grid_widths) <= 0:
            issues.append(f"table {index}: empty table grid")
        for row_index, row in enumerate(table.findall("./w:tr", NS), start=1):
            if row.find("./w:trPr/w:cantSplit", NS) is None:
                issues.append(f"table {index} row {row_index}: split protection missing")
                break
            cells = row.findall("./w:tc", NS)
            for cell_index, cell in enumerate(cells, start=1):
                tc_w = cell.find("./w:tcPr/w:tcW", NS)
                if tc_w is None or tc_w.attrib.get(qn("w:type")) != "dxa":
                    issues.append(f"table {index} cell {cell_index}: non-DXA cell width")
                    break
    return issues


def audit_docx(docx: Path, source: Path) -> dict[str, object]:
    issues: list[str] = []
    with zipfile.ZipFile(docx) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml")
        xml_text = document_xml.decode("utf-8")
        text = document_visible_text(zf)
        media = [name for name in names if name.startswith("word/media/")]
        hyperlinks = relationship_targets(zf)
        if "[[TOC]]" in xml_text or "```" in xml_text or "~~~" in xml_text:
            issues.append("Markdown placeholder/fence leaked")
        if "图像资源未找到" in text:
            issues.append("missing-image warning present")
        if "wp:anchor" in xml_text:
            issues.append("floating image anchor present")
        if xml_text.count("wp:inline") < 1:
            issues.append("cover inline image missing")
        issues.extend(table_geometry_issues(document_xml))

        core = ET.fromstring(zf.read("docProps/core.xml"))
        creator = core.find("dc:creator", NS)
        modified = core.find("cp:lastModifiedBy", NS)
        if creator is not None and (creator.text or "").strip().lower() in {"why", "administrator", "admin"}:
            issues.append("personal creator metadata present")
        if modified is not None and (modified.text or "").strip().lower() in {"why", "administrator", "admin"}:
            issues.append("personal lastModifiedBy metadata present")

        for target in hyperlinks:
            if target.lower().endswith(".docx") and "/" not in target and "\\" not in target:
                if not (OUT / target).exists():
                    issues.append(f"broken sibling DOCX link: {target}")

    doc = Document(str(docx))
    if len(doc.sections) < 2:
        issues.append("cover/body section split missing")
    for idx, section in enumerate(doc.sections):
        if abs(section.page_width.mm - builder.PAGE_W_MM) > 0.6:
            issues.append(f"section {idx}: page width is not A4")
        if abs(section.page_height.mm - builder.PAGE_H_MM) > 0.6:
            issues.append(f"section {idx}: page height is not A4")

    snippets = source_snippets(source)
    doc_norm = norm(text)
    missing = [snippet for snippet in snippets if norm(snippet) not in doc_norm]
    coverage = 1.0 - len(missing) / max(1, len(snippets))
    if coverage < 0.985:
        issues.append(f"source-line fidelity below threshold: {coverage:.3%}")

    return {
        "file": docx.name,
        "bytes": docx.stat().st_size,
        "sections": len(doc.sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "media": len(media),
        "hyperlinks": len(hyperlinks),
        "source_snippets": len(snippets),
        "missing_snippets": len(missing),
        "fidelity": round(coverage, 6),
        "missing_examples": missing[:12],
        "issues": issues,
    }


def natural_page_key(path: Path) -> int:
    match = re.search(r"(\d+)$", path.stem)
    return int(match.group(1)) if match else 0


def audit_render(code: str) -> tuple[dict[str, object], list[Path]]:
    rerendered = RERENDERS / code
    directory = rerendered if rerendered.exists() else RENDERS / code
    pages = sorted(directory.glob("page-*.png"), key=natural_page_key)
    issues: list[str] = []
    records = []
    dimensions = set()
    for index, path in enumerate(pages, start=1):
        with Image.open(path) as im:
            rgb = np.asarray(im.convert("RGB"), dtype=np.int16)
        dimensions.add((rgb.shape[1], rgb.shape[0]))
        delta = np.max(np.abs(255 - rgb), axis=2)
        mask = delta > 14
        ratio = float(mask.mean())
        ys, xs = np.where(mask)
        bbox = None if len(xs) == 0 else [int(xs.min()), int(ys.min()), int(xs.max()), int(ys.max())]
        edge = bool(mask[:3, :].any() or mask[-3:, :].any() or mask[:, :3].any() or mask[:, -3:].any())
        blank = ratio < 0.0025
        if edge:
            issues.append(f"page {index}: non-white pixels touch physical page edge")
        if blank:
            issues.append(f"page {index}: suspiciously blank ({ratio:.4f})")
        records.append({"page": index, "ink_ratio": round(ratio, 5), "bbox": bbox, "edge_touch": edge, "blank": blank})
    if not pages:
        issues.append("no rendered pages")
    if len(dimensions) > 1:
        issues.append(f"mixed raster dimensions: {sorted(dimensions)}")
    return {
        "code": code,
        "pages": len(pages),
        "dimensions": sorted(dimensions),
        "min_ink_ratio": min((r["ink_ratio"] for r in records), default=0),
        "max_ink_ratio": max((r["ink_ratio"] for r in records), default=0),
        "issues": issues,
        "page_metrics": records,
    }, pages


def contact_sheet(code: str, pages: list[Path], start: int, batch: int = 20) -> Path:
    CONTACTS.mkdir(parents=True, exist_ok=True)
    subset = pages[start : start + batch]
    cols, rows = 5, 4
    thumb_w, thumb_h = 220, 311
    gutter_x, gutter_y = 22, 38
    top = 48
    width = cols * thumb_w + (cols + 1) * gutter_x
    height = top + rows * thumb_h + (rows + 1) * gutter_y
    sheet = Image.new("RGB", (width, height), (236, 240, 244))
    draw = ImageDraw.Draw(sheet)
    try:
        font = ImageFont.truetype(str(builder.FONT_BOLD), 20)
        label_font = ImageFont.truetype(str(builder.FONT_REGULAR), 16)
    except Exception:
        font = ImageFont.load_default()
        label_font = ImageFont.load_default()
    draw.text((gutter_x, 12), f"DOC {code}  |  PAGES {start + 1}-{start + len(subset)}", font=font, fill=(16, 36, 62))
    for offset, page in enumerate(subset):
        row, col = divmod(offset, cols)
        x = gutter_x + col * (thumb_w + gutter_x)
        y = top + gutter_y + row * (thumb_h + gutter_y)
        with Image.open(page) as im:
            thumb = ImageOps.contain(im.convert("RGB"), (thumb_w, thumb_h), Image.Resampling.LANCZOS)
        sheet.paste(thumb, (x, y))
        draw.rectangle((x - 1, y - 1, x + thumb.width, y + thumb.height), outline=(163, 175, 187), width=1)
        draw.text((x, y - 22), f"{code}-{start + offset + 1:03d}", font=label_font, fill=(64, 78, 92))
    out = CONTACTS / f"{code}-{start + 1:03d}-{start + len(subset):03d}.jpg"
    sheet.save(out, quality=92, subsampling=0)
    return out


def main() -> int:
    docs = sorted(OUT.glob("[0-1][0-9]-*.docx"))
    expected = sorted(SRC.glob("[0-1][0-9]-*.md"))
    expected_codes = [path.stem[:2] for path in expected if path.stem[:2] in builder.DOC_INFO]
    found_codes = [path.stem[:2] for path in docs]
    top_issues = []
    if found_codes != expected_codes:
        top_issues.append(f"DOCX set mismatch: expected {expected_codes}, found {found_codes}")

    doc_audits = []
    render_audits = []
    contact_sheets = []
    for docx in docs:
        code = docx.stem[:2]
        source_matches = list(SRC.glob(f"{code}-*.md"))
        if not source_matches:
            top_issues.append(f"missing source for {docx.name}")
            continue
        doc_audits.append(audit_docx(docx, source_matches[0]))
        render, pages = audit_render(code)
        render_audits.append(render)
        for start in range(0, len(pages), 20):
            contact_sheets.append(str(contact_sheet(code, pages, start).relative_to(ROOT)))

    report = {
        "documents": len(docs),
        "total_pages": sum(item["pages"] for item in render_audits),
        "top_issues": top_issues,
        "docx_audits": doc_audits,
        "render_audits": render_audits,
        "contact_sheets": contact_sheets,
    }
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    summary = {
        "documents": report["documents"],
        "total_pages": report["total_pages"],
        "top_issues": top_issues,
        "docx_issues": {item["file"]: item["issues"] for item in doc_audits if item["issues"]},
        "render_issues": {item["code"]: item["issues"] for item in render_audits if item["issues"]},
        "fidelity": {item["file"]: item["fidelity"] for item in doc_audits},
        "contact_sheets": len(contact_sheets),
        "report": str(REPORT_PATH),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    has_issues = bool(top_issues or any(item["issues"] for item in doc_audits) or any(item["issues"] for item in render_audits))
    return 1 if has_issues else 0


if __name__ == "__main__":
    raise SystemExit(main())
